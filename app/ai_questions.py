"""
Модуль для работы с AI-генерацией вопросов по инструктажам
Использует OpenAI GPT-4o-mini для генерации вопросов на русском и казахском языках
"""

import os
from pathlib import Path
from typing import Dict, List
import json
import pypdf
import docx
from openai import OpenAI

def extract_text_from_pdf(file_path: str) -> str:
    """Извлекает текст из PDF файла"""
    try:
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Извлекает текст из DOCX файла"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""


def extract_text_from_file(file_path: str) -> str:
    """Универсальная функция извлечения текста из файла"""
    ext = Path(file_path).suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def generate_questions_openai(
    text: str, 
    api_key: str, 
    language: str = 'ru',
    num_questions: int = 4
) -> Dict:
    """
    Генерирует вопросы через OpenAI GPT-4o-mini
    
    Args:
        text: Текст инструктажа
        api_key: OpenAI API ключ
        language: Язык вопросов ('ru' или 'kk')
        num_questions: Количество вопросов (по умолчанию 4)
    
    Returns:
        {
            "questions": [
                {
                    "question": "Текст вопроса?",
                    "options": ["A) Вариант 1", "B) Вариант 2", "C) Вариант 3", "D) Вариант 4"],
                    "correct": "A",
                    "explanation": "Пояснение"
                },
                ...
            ]
        }
    """
    client = OpenAI(api_key=api_key)
    
    # Ограничиваем текст (для экономии токенов)
    text_truncated = text[:6000] if len(text) > 6000 else text
    
    if language == 'kk':
        # Промпт на казахском
        prompt = f"""Сіз еңбек қауіпсіздігі бойынша сарапшысыз. Осы нұсқаулықты талдап, түсінуді тексеруге арналған {num_questions} сұрақ жасаңыз.

ӨТЕ МАҢЫЗДЫ:
1. Сұрақтар ҚАЗАҚ ТІЛІНДЕ болуы керек
2. Әр сұрақта 4 жауап нұсқасы болуы керек (A, B, C, D)
3. Тек БІР дұрыс жауап
4. ДҰРЫС ЕМЕС нұсқалар ШЫНАЙЫ және НАҚТЫ болуы керек
5. Анық қате нұсқаларды қолданбаңыз
6. Дұрыс емес жауаптар нақты, бірақ толық емес немесе дәл емес ережелердің түсіндірмесі сияқты естілуі керек
7. Сұрақтар нұсқаулықтан алынған НАҚТЫ талаптарды ДӘЛМЕ-ДӘЛ түсінуді тексеруі керек

ЖАУАП ФОРМАТЫ (тек JSON):
{{
  "questions": [
    {{
      "question": "Нұсқаулықтан нақты сұрақ?",
      "options": [
        "A) Шынайы бірақ толық емес нұсқа",
        "B) Мәтіннен дұрыс жауап",
        "C) Басқа шынайы нұсқа",
        "D) Тағы бір нақты нұсқа"
      ],
      "correct": "B",
      "explanation": "Мәтінге сілтеме жасай отырып қысқаша түсініктеме"
    }}
  ]
}}

НҰСҚАУЛЫҚ МӘТІНІ:
{text_truncated}

Тек таза JSON қайтарыңыз."""
    else:
        # Промпт на русском
        prompt = f"""Ты эксперт по технике безопасности. Проанализируй этот инструктаж и создай {num_questions} вопроса для проверки понимания.

КРИТИЧЕСКИ ВАЖНО:
1. Вопросы на РУССКОМ языке (даже если текст на казахском)
2. Каждый вопрос с 4 вариантами ответа (A, B, C, D)
3. Только ОДИН правильный ответ
4. НЕПРАВИЛЬНЫЕ варианты должны быть ПРАВДОПОДОБНЫМИ и реалистичными
5. НЕ используй очевидно глупые варианты типа "игнорировать инструкции"
6. Неправильные ответы должны звучать как реальные, но неполные или неточные интерпретации правил
7. Вопросы должны проверять ТОЧНОЕ понимание конкретных требований из инструктажа

ФОРМАТ ОТВЕТА (строго JSON):
{{
  "questions": [
    {{
      "question": "Конкретный вопрос из инструктажа?",
      "options": [
        "A) Правдоподобный но неполный вариант",
        "B) Правильный ответ из текста",
        "C) Другой правдоподобный вариант",
        "D) Еще один реалистичный вариант"
      ],
      "correct": "B",
      "explanation": "Краткое пояснение со ссылкой на текст"
    }}
  ]
}}

ТЕКСТ ИНСТРУКТАЖА:
{text_truncated}

Верни только валидный JSON без markdown и комментариев."""

    try:
        system_message = "Сіз қауіпсіздік бойынша сұрақтар жасау көмекшісісіз. Әрқашан дұрыс JSON қайтарыңыз." if language == 'kk' else "Ты помощник для создания вопросов по технике безопасности. Всегда отвечай валидным JSON."
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        
        content = response.choices[0].message.content.strip()
        
        # Убираем markdown разметку если есть
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        result = json.loads(content)
        return result
        
    except json.JSONDecodeError as e:
        raise RuntimeError(f"OpenAI вернул невалидный JSON: {e}. Ответ: {content[:300]}")
    except Exception as e:
        raise RuntimeError(f"Ошибка генерации вопросов: {str(e)}")


def validate_answers(questions: List[Dict], user_answers: Dict[str, str]) -> Dict:
    """
    Проверяет ответы пользователя
    
    Args:
        questions: Список вопросов с правильными ответами
        user_answers: Словарь ответов пользователя {"0": "B", "1": "A", ...}
    
    Returns:
        {
            "correct_count": 3,
            "total_count": 4,
            "score_percentage": 75.0,
            "passed": True,
            "details": [
                {"question_index": 0, "correct": True, "user_answer": "B", "correct_answer": "B"},
                ...
            ]
        }
    """
    correct_count = 0
    details = []
    
    for i, question in enumerate(questions):
        user_answer = user_answers.get(str(i), "")
        correct_answer = question["correct"]
        is_correct = (user_answer == correct_answer)
        
        if is_correct:
            correct_count += 1
        
        details.append({
            "question_index": i,
            "correct": is_correct,
            "user_answer": user_answer,
            "correct_answer": correct_answer
        })
    
    total_count = len(questions)
    score_percentage = (correct_count / total_count * 100) if total_count > 0 else 0
    passed = score_percentage >= 75.0  # Минимум 75% для прохождения
    
    return {
        "correct_count": correct_count,
        "total_count": total_count,
        "score_percentage": round(score_percentage, 2),
        "passed": passed,
        "details": details
    }
